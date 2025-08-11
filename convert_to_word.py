#!/usr/bin/env python3
"""
Convert CloudInit Status Report from Markdown to Word format
"""

import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new run object and add the text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color and underline for hyperlink styling
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0563C1")
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def parse_markdown_to_word():
    """
    Parse the Markdown status report and convert to Word format
    """
    # Read the markdown file
    with open('cloudinit_status_report.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Handle main title (# heading)
        if line.startswith('# '):
            title = line[2:].strip()
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Handle section headers (## heading)
        elif line.startswith('## '):
            header = line[3:].strip()
            doc.add_heading(header, level=2)
            
        # Handle subsection headers (### heading)
        elif line.startswith('### '):
            subheader = line[4:].strip()
            doc.add_heading(subheader, level=3)
            
        # Handle italic text (*text*)
        elif line.startswith('*') and line.endswith('*') and len(line) > 2:
            italic_text = line[1:-1]
            para = doc.add_paragraph()
            run = para.add_run(italic_text)
            run.italic = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Handle bold text (**text**)
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    para.add_run(part)
                else:
                    run = para.add_run(part)
                    run.bold = True
                    
        # Handle bullet points (- text)
        elif line.startswith('- '):
            bullet_text = line[2:].strip()
            # Check if it contains bold formatting
            if '**' in bullet_text:
                para = doc.add_paragraph(style='List Bullet')
                parts = bullet_text.split('**')
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        para.add_run(part)
                    else:
                        run = para.add_run(part)
                        run.bold = True
            else:
                doc.add_paragraph(bullet_text, style='List Bullet')
                
        # Handle horizontal rules (---)
        elif line.startswith('---'):
            para = doc.add_paragraph()
            para.add_run('_' * 80)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Handle regular paragraphs
        elif line and not line.startswith('#'):
            # Check for bold formatting in regular text
            if '**' in line:
                para = doc.add_paragraph()
                parts = line.split('**')
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        para.add_run(part)
                    else:
                        run = para.add_run(part)
                        run.bold = True
            else:
                doc.add_paragraph(line)
        
        i += 1
    
    # Add footer with generation info
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Document generated on: ').bold = True
    footer_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    footer_para.add_run('\nSource: Azure DevOps CloudInit Epic Status Report').italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    output_filename = f'CloudInit_Status_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(output_filename)
    
    print(f"✅ Word document created successfully: {output_filename}")
    return output_filename

if __name__ == "__main__":
    try:
        output_file = parse_markdown_to_word()
        print(f"\n📄 Report generated: {output_file}")
        print("📁 Location: Current directory")
        print("✉️  Ready to share with your team!")
        
    except FileNotFoundError:
        print("❌ Error: cloudinit_status_report.md file not found")
        print("Please ensure the markdown report exists in the current directory")
        
    except Exception as e:
        print(f"❌ Error generating Word document: {str(e)}")
        print("Please ensure you have the required packages installed:")
        print("pip install python-docx")
