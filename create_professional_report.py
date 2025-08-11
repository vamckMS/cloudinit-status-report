#!/usr/bin/env python3
"""
Enhanced CloudInit Status Report to Word Converter
Creates a professionally formatted Word document from the markdown report
"""

import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def create_professional_word_report():
    """
    Create a professional Word document from the CloudInit status report
    """
    # Read the markdown file
    try:
        with open('cloudinit_status_report.md', 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print("❌ Error: cloudinit_status_report.md not found")
        return None
    
    # Create new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Set default font to Aptos for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Aptos'
    font.size = Pt(11)
    
    # Customize heading styles
    # Heading 1 style
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Aptos'
    heading1_style.font.size = Pt(20)
    heading1_style.font.bold = True
    
    # Heading 2 style
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Aptos'
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True
    
    # Heading 3 style - increased size as requested
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.name = 'Aptos'
    heading3_style.font.size = Pt(14)  # Increased from default ~13pt
    heading3_style.font.bold = True
    
    # Customize List Bullet style for tighter spacing
    try:
        list_style = doc.styles['List Bullet']
        list_style.font.name = 'Aptos'
        list_style.font.size = Pt(11)
        # Reduce paragraph spacing
        list_style.paragraph_format.space_before = Pt(2)
        list_style.paragraph_format.space_after = Pt(2)
        list_style.paragraph_format.line_spacing = 1.1
    except KeyError:
        # Create custom list style if default doesn't exist
        pass
    
    # Parse content
    lines = content.split('\n')
    
    # Add title
    title = doc.add_heading('CloudInit Epic Status Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation info
    gen_para = doc.add_paragraph()
    gen_para.add_run('Generated on: ').bold = True
    gen_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Apply Aptos font to generation info
    for run in gen_para.runs:
        run.font.name = 'Aptos'
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Empty line
    
    # Project info section
    doc.add_heading('Project Information', level=2)
    project_info = doc.add_paragraph()
    project_info.add_run('Project: ').bold = True
    project_info.add_run('One')
    project_info.add_run('\nArea: ').bold = True
    project_info.add_run('One\\CloudInit')
    project_info.add_run('\nIteration: ').bold = True
    project_info.add_run('One\\Bromine')
    # Apply Aptos font to project info
    for run in project_info.runs:
        run.font.name = 'Aptos'
        run.font.size = Pt(11)
    # Tighter spacing for this section
    project_info.paragraph_format.space_before = Pt(3)
    project_info.paragraph_format.space_after = Pt(6)
    project_info.paragraph_format.line_spacing = 1.1
    
    # Summary section
    doc.add_heading('Summary', level=2)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    # Add summary data
    summary_data = [
        ['Total Epics in CloudInit/Bromine', '10'],
        ['Active Epics', '5'],
        ['New Epics', '5'],
        ['Completed Epics', '0'],
        ['Total Child Work Items', '25+'],
        ['- Done', '12'],
        ['- Active/In Progress', '7'],
        ['- To Do/New', '6+']
    ]
    
    for row_data in summary_data:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        # Make first column bold
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Epic Details section
    doc.add_heading('Epic Details', level=2)
    
    # Parse epics from markdown
    epic_pattern = r'### Epic \d+: (.+)'
    id_pattern = r'\*\*ID\*\*: (\d+)'
    state_pattern = r'\*\*State\*\*: (\w+)'
    assignee_pattern = r'\*\*Assignee\*\*: (.+)'
    
    epic_sections = content.split('### Epic')[1:]  # Skip the part before first epic
    
    for i, epic_section in enumerate(epic_sections):
        lines = epic_section.strip().split('\n')
        if not lines:
            continue
            
        # Extract epic title
        title_match = re.match(r'\d+: (.+)', lines[0])
        if title_match:
            epic_title = title_match.group(1)
        else:
            epic_title = lines[0]
        
        # Add epic as heading
        doc.add_heading(f'Epic {i+1}: {epic_title}', level=3)
        
        # Extract epic details
        epic_details = []
        child_items = []
        in_child_section = False
        
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('**Child Items**:'):
                in_child_section = True
                continue
            elif line.startswith('---'):
                break
            elif in_child_section:
                if line.startswith('- **'):
                    child_items.append(line[2:])  # Remove '- '
            else:
                # Clean up markdown formatting for basic details
                clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
                if clean_line.startswith('- **') and clean_line.endswith('**'):
                    epic_details.append(clean_line[2:])  # Remove bullet
                elif clean_line and not clean_line.startswith('#'):
                    epic_details.append(clean_line)
        
        # Add epic details
        for detail in epic_details:
            if ':' in detail:
                para = doc.add_paragraph()
                parts = detail.split(':', 1)
                bold_run = para.add_run(parts[0] + ': ')
                bold_run.bold = True
                bold_run.font.name = 'Aptos'
                bold_run.font.size = Pt(11)
                
                normal_run = para.add_run(parts[1].strip())
                normal_run.font.name = 'Aptos'
                normal_run.font.size = Pt(11)
                
                # Tighter spacing for epic details
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.1
            else:
                para = doc.add_paragraph(detail)
                for run in para.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(11)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
        
        # Add child items if any
        if child_items:
            child_header = doc.add_paragraph()
            child_header_run = child_header.add_run('Child Items:')
            child_header_run.bold = True
            child_header_run.font.name = 'Aptos'
            child_header_run.font.size = Pt(11)
            child_header.paragraph_format.space_before = Pt(4)
            child_header.paragraph_format.space_after = Pt(2)
            
            for item in child_items:
                # Clean up markdown formatting
                clean_item = re.sub(r'\*\*(\d+)\*\*', r'\1', item)
                clean_item = re.sub(r'\*([^*]+)\*', r'\1', clean_item)
                
                bullet_para = doc.add_paragraph(clean_item, style='List Bullet')
                for run in bullet_para.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)  # Slightly smaller for child items
                # Extra tight spacing for bullet items
                bullet_para.paragraph_format.space_before = Pt(1)
                bullet_para.paragraph_format.space_after = Pt(1)
                bullet_para.paragraph_format.line_spacing = 1.0
                bullet_para.paragraph_format.left_indent = Inches(0.25)
        else:
            para = doc.add_paragraph()
            header_run = para.add_run('Child Items: ')
            header_run.bold = True
            header_run.font.name = 'Aptos'
            header_run.font.size = Pt(11)
            
            no_items_run = para.add_run('No child items found')
            no_items_run.font.name = 'Aptos'
            no_items_run.font.size = Pt(11)
            
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(2)
        
        doc.add_paragraph()  # Empty line between epics
    
    # Key Observations section
    doc.add_heading('Key Observations', level=2)
    
    observations = [
        ('Active Development Areas:', [
            'Linux Provisioning Analyzer (LPA): Multiple epics focused on chatbot development and quality improvements',
            'Aurora Migration: Testing and migration work in progress',
            'Security Enhancements: LPA security improvements for Bromine',
            'Infrastructure: NVMe naming utilities and container support'
        ]),
        ('Recent Activity:', [
            'Three new LPA chatbot versions (V3.5, V4, V4.5) were created on July 31, 2025',
            'Most epics have been updated within the last 2 weeks',
            'Strong focus on Linux Provisioning Analyzer ecosystem'
        ]),
        ('Team Distribution:', [
            'Christopher Patterson: 2 epics (NVMe naming, Quality improvements)',
            'Peyton Robertson: 2 epics (Azure-init containers, LPA Security)',
            'Sophie Gee: 1 epic (LPA Quality improvements)',
            'Ben Ryan: 1 epic (Aurora migration)',
            'Vincenzo Marcella: 1 epic (LPA Chatbot V3)',
            'Unassigned: 3 epics (LPA Chatbot V3.5, V4, V4.5)'
        ])
    ]
    
    for section_title, items in observations:
        doc.add_heading(section_title, level=3)
        for item in items:
            bullet_para = doc.add_paragraph(item, style='List Bullet')
            for run in bullet_para.runs:
                run.font.name = 'Aptos'
                run.font.size = Pt(11)
            # Tighter spacing for observation bullets
            bullet_para.paragraph_format.space_before = Pt(1)
            bullet_para.paragraph_format.space_after = Pt(1)
            bullet_para.paragraph_format.line_spacing = 1.1
            bullet_para.paragraph_format.left_indent = Inches(0.25)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run1 = footer.add_run('Report generated using Azure DevOps MCP integration on ')
    footer_run1.italic = True
    footer_run1.font.name = 'Aptos'
    footer_run1.font.size = Pt(10)
    
    footer_run2 = footer.add_run(datetime.now().strftime('%Y-%m-%d'))
    footer_run2.italic = True
    footer_run2.font.name = 'Aptos'
    footer_run2.font.size = Pt(10)
    
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f'CloudInit_Epic_Status_Report_Enhanced_{timestamp}.docx'
    doc.save(output_filename)
    
    # Also save a version with a simple name for easy access
    simple_filename = 'CloudInit_Status_Report_Latest.docx'
    doc.save(simple_filename)
    
    return output_filename

if __name__ == "__main__":
    try:
        output_file = create_professional_word_report()
        if output_file:
            print(f"✅ Enhanced Word document created: {output_file}")
            print(f"📁 Also saved as: CloudInit_Status_Report_Latest.docx")
            print("📊 Enhanced formatting features:")
            print("   • Aptos font throughout document")
            print("   • Larger Heading 3 size (14pt)")
            print("   • Tighter bullet spacing and indentation")
            print("   • Professional layout with improved readability")
            print("   • Summary table and structured sections")
            print("✉️  Ready to share with stakeholders!")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        print("Please ensure python-docx is installed: pip install python-docx")
